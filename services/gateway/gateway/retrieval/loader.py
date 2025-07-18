"""Document loading from files and remote APIs."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from typing import TYPE_CHECKING

from langchain_core.documents import Document
from loguru import logger

if TYPE_CHECKING:
    from gateway.clients.data_api import (
        DataServiceClient,
    )


class CorpusLoader:
    """Load documents from various sources."""

    def fetch_from_api(
        self,
        client: DataServiceClient,
    ) -> list[Document]:
        """Download corpus archive and extract."""
        try:
            archive_bytes = client.download_corpus_archive()
        except Exception:
            logger.warning(
                "Could not download corpus archive",
            )
            return []

        documents: list[Document] = []
        with zipfile.ZipFile(
            io.BytesIO(archive_bytes),
        ) as zf:
            for name in zf.namelist():
                if not name.endswith(".txt"):
                    continue
                content = zf.read(name).decode(
                    "utf-8",
                    errors="replace",
                )
                if content.strip():
                    documents.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": name,
                                "origin": "api",
                            },
                        ),
                    )
        logger.info(
            "Loaded {} documents from API archive",
            len(documents),
        )
        return documents

    def scan_directory(
        self,
        path: str | Path,
    ) -> list[Document]:
        """Read all supported files from a dir."""
        directory = Path(path)
        if not directory.is_dir():
            logger.warning(
                "Directory not found: {}",
                path,
            )
            return []

        extensions = {
            ".txt",
            ".md",
            ".pdf",
            ".docx",
        }
        documents: list[Document] = []

        for file_path in sorted(
            directory.rglob("*"),
        ):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                content = self.parse_file(file_path)
                if content.strip():
                    documents.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": str(
                                    file_path,
                                ),
                                "origin": "local",
                            },
                        ),
                    )

        logger.info(
            "Scanned {} documents from {}",
            len(documents),
            path,
        )
        return documents

    def parse_file(
        self,
        path: str | Path,
    ) -> str:
        """Extract text from a single file."""
        file_path = Path(path)
        suffix = file_path.suffix.lower()

        if suffix in (".txt", ".md"):
            return file_path.read_text(
                encoding="utf-8",
                errors="replace",
            )

        if suffix == ".pdf":
            return self._read_pdf(file_path)

        if suffix == ".docx":
            return self._read_docx(file_path)

        logger.warning(
            "Unsupported file type: {}",
            suffix,
        )
        return ""

    @staticmethod
    def _read_pdf(path: Path) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except Exception:
            logger.error(
                "Failed to read PDF: {}",
                path,
            )
            return ""

    @staticmethod
    def _read_docx(path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDoc

            doc = DocxDoc(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception:
            logger.error(
                "Failed to read DOCX: {}",
                path,
            )
            return ""
